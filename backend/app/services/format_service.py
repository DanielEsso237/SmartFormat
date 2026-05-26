import re
import copy
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

from app.charte.forms import get_filiere_color
from app.schemas.document import DocumentAnalysis

FONT_NAME    = "Times New Roman"
FONT_BODY    = 12
FONT_H1      = 14
FONT_H2      = 13
FONT_H3      = 12
LINE_SPACING = 1.5
MARGIN_CM    = 2.5

LOGO_PATH = Path(__file__).resolve().parent.parent / "assets" / "logo.png"

TOC_STYLE_IDS = {"toc1","toc2","toc3","toc4","tocheading","tableofcontents"}
TOC_FIELD_RE  = re.compile(r'\bTOC\b')

HEADING_SCORE_THRESHOLD = 4

COVER_STOP_RE = re.compile(
    r'^(introduction|r[eé]sum[eé]|abstract|remerciements?|d[eé]dicace'
    r'|sommaire|table\s+des\s+mati'
    r'|i[\-\.\s]|ii[\-\.\s]|iii[\-\.\s]'
    r'|chapitre\s+\d|partie\s+\d)',
    re.IGNORECASE
)

HEADING_FALSE_POSITIVE_RE = re.compile(
    r'(comprend|suivants?|notamment|ainsi\s+que|dont|suivante|:$|\.\s*$)',
    re.IGNORECASE
)

PRELIM_SECTIONS = [
    ('dedicace',      re.compile(r'^d[eé]dicace',               re.I)),
    ('remerciements', re.compile(r'^remerciements?',            re.I)),
    ('resume',        re.compile(r'^r[eé]sum[eé]',              re.I)),
    ('abstract',      re.compile(r'^abstract',                  re.I)),
    ('sommaire',      re.compile(r'^(sommaire|table\s+des\s+mati)', re.I)),
    ('introduction',  re.compile(r'^introduction',              re.I)),
]

PRELIM_LABELS = {
    'dedicace':      'DÉDICACE',
    'remerciements': 'REMERCIEMENTS',
    'resume':        'RÉSUMÉ',
    'abstract':      'ABSTRACT',
    'sommaire':      'TABLE DES MATIÈRES',
    'introduction':  'INTRODUCTION',
}

REQUIRED_PRELIMS = {
    'memoire':        ['dedicace','remerciements','resume','abstract','sommaire','introduction'],
    'rapport_stage':  ['dedicace','remerciements','resume','abstract','sommaire','introduction'],
    'rapport_projet': ['dedicace','remerciements','resume','abstract','sommaire','introduction'],
    'expose':         ['sommaire','introduction'],
    'demande':        [],
}


def _set_margins(section, cm=MARGIN_CM):
    section.top_margin    = Cm(cm)
    section.bottom_margin = Cm(cm)
    section.left_margin   = Cm(cm)
    section.right_margin  = Cm(cm)


def _apply_base_style(doc):
    s = doc.styles['Normal']
    s.font.name = FONT_NAME
    s.font.size = Pt(FONT_BODY)
    s.paragraph_format.line_spacing      = LINE_SPACING
    s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    s.paragraph_format.space_after       = Pt(6)
    s.paragraph_format.space_before      = Pt(0)


def _para(doc, text="", size=FONT_BODY, bold=False, italic=False,
          align="justify", space_before=0, space_after=6,
          color=None, keep_with_next=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before   = Pt(space_before)
    p.paragraph_format.space_after    = Pt(space_after)
    p.paragraph_format.line_spacing   = LINE_SPACING
    p.paragraph_format.keep_with_next = keep_with_next
    p.alignment = {
        "left":    WD_ALIGN_PARAGRAPH.LEFT,
        "center":  WD_ALIGN_PARAGRAPH.CENTER,
        "right":   WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    if text:
        r = p.add_run(text)
        r.font.name   = FONT_NAME
        r.font.size   = Pt(size)
        r.font.bold   = bold
        r.font.italic = italic
        if color:
            r.font.color.rgb = RGBColor(*color)
    return p


def _heading(doc, text, level=1):
    cfg = {
        1: dict(size=FONT_H1, align="center", space_before=12, space_after=6),
        2: dict(size=FONT_H2, align="left",   space_before=10, space_after=4),
        3: dict(size=FONT_H3, align="left",   space_before=8,  space_after=4),
    }.get(level, dict(size=FONT_H2, align="left", space_before=10, space_after=4))
    display = text.upper() if level == 1 else text
    _para(doc, display, bold=True, keep_with_next=True, **cfg)


def _page_break(doc):
    doc.add_page_break()


def _is_heading(para) -> tuple[bool, int]:
    text = para.text.strip()
    if not text or len(text) < 2 or len(text) > 150:
        return False, 0
    if HEADING_FALSE_POSITIVE_RE.search(text):
        return False, 0

    style_name = (para.style.name or "").lower()
    if "heading" in style_name:
        m = re.search(r'(\d)', style_name)
        return True, min(int(m.group(1)) if m else 1, 3)

    sizes = [r.font.size for r in para.runs if r.text.strip() and r.font.size]
    size  = round(sizes[0] / 12700) if sizes else None
    bolds = [r.bold for r in para.runs if r.text.strip()]
    is_bold  = any(b for b in bolds if b is not None)
    is_upper = text == text.upper() and len(text) > 3 and bool(re.search(r'[A-Z]', text))

    roman_h1 = bool(re.match(r'^(I|II|III|IV|V|VI|VII|VIII|IX|X)[\-\.\s]', text))
    roman_h2 = bool(re.match(r'^(i|ii|iii|iv|v|vi|vii|viii|ix|x)[\-\.\s]', text))
    num_h2   = bool(re.match(r'^\d+[\-\.\)\s]', text))
    alpha_h2 = bool(re.match(r'^[A-Z][\-\.\)\s]', text))

    score = 0
    if size and size >= 14: score += 3
    elif size and size == 13: score += 2
    if is_bold:  score += 2
    if is_upper: score += 2
    if roman_h1: score += 3
    if num_h2 or alpha_h2 or roman_h2: score += 2

    if score < HEADING_SCORE_THRESHOLD:
        return False, 0

    if roman_h1 or (is_upper and size and size >= 13):
        return True, 1
    if num_h2 or alpha_h2 or roman_h2:
        return True, 2
    if is_upper or (size and size >= 14):
        return True, 1
    return True, 2


def _detect_cover_end(paragraphs) -> int:
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if not text:
            continue
        if COVER_STOP_RE.match(text):
            return i
        if i > 3 and len(text) > 100:
            is_h, _ = _is_heading(p)
            if not is_h:
                return i
    return min(8, len(paragraphs))


def _scan_present_sections(paragraphs, cover_end) -> dict[str, list]:
    """
    Parcourt le source depuis cover_end.
    Retourne un dict: clé_section → [paragraphes (text) du contenu de cette section]
    La clé 'body' contient tout ce qui est hors sections préliminaires.
    """
    result = {k: [] for k, _ in PRELIM_SECTIONS}
    result['body'] = []

    current_section = None
    prelim_keys = [k for k, _ in PRELIM_SECTIONS]

    for i, p in enumerate(paragraphs):
        if i < cover_end:
            continue
        text = p.text.strip()
        if not text:
            continue

        matched = None
        for key, pat in PRELIM_SECTIONS:
            if pat.match(text):
                matched = key
                break

        if matched:
            current_section = matched
            continue

        if current_section and current_section in prelim_keys:
            is_h, _ = _is_heading(p)
            if is_h:
                current_section = 'body'
            else:
                result[current_section].append(p)
                continue

        current_section = 'body'
        result['body'].append(p)

    return result


def _extract_toc(source_path: Path) -> dict:
    if not source_path or not source_path.exists():
        return {'type': None, 'elements': []}
    try:
        src = Document(source_path)
    except Exception:
        return {'type': None, 'elements': []}

    body = src.element.body

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'sdt':
            for fld in child.iter(qn('w:instrText')):
                if fld.text and TOC_FIELD_RE.search(fld.text):
                    return {'type': 'sdt', 'elements': [copy.deepcopy(child)]}

    toc_paras, in_toc = [], False
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'p':
            if in_toc: break
            continue
        has_toc = any(TOC_FIELD_RE.search(f.text or '') for f in child.iter(qn('w:instrText')))
        if has_toc:
            in_toc = True
        if in_toc:
            toc_paras.append(copy.deepcopy(child))
            for fc in child.iter(qn('w:fldChar')):
                if fc.get(qn('w:fldCharType')) == 'end' and not has_toc:
                    in_toc = False
    if toc_paras:
        return {'type': 'fldchar', 'elements': toc_paras}

    toc_style_paras = []
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'p': continue
        ps = child.find('.//' + qn('w:pStyle'))
        if ps is not None:
            sid = ps.get(qn('w:val'), '').lower().replace(' ', '').replace('-', '')
            if sid in TOC_STYLE_IDS or sid.startswith('toc'):
                toc_style_paras.append(copy.deepcopy(child))
    if toc_style_paras:
        return {'type': 'styles', 'elements': toc_style_paras}

    return {'type': None, 'elements': []}


def _insert_toc(doc, toc):
    for el in toc['elements']:
        doc.element.body.append(el)


def _write_section_content(doc, paragraphs):
    """Écrit une liste de paragraphes source avec la charte appliquée."""
    for p in paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if re.match(r'^[-•–]\s+', text):
            body_p = doc.add_paragraph(style="List Bullet")
            body_p.paragraph_format.line_spacing = LINE_SPACING
            body_p.paragraph_format.space_after  = Pt(4)
            r = body_p.add_run(re.sub(r'^[-•–]\s+', '', text))
            r.font.name = FONT_NAME
            r.font.size = Pt(FONT_BODY)
        else:
            _para(doc, text, size=FONT_BODY, align="justify",
                  space_before=0, space_after=6)


def _write_prelim_section(doc, key, content_paras, toc, doc_type):
    """
    Écrit une section préliminaire sur sa propre page.
    Si content_paras est vide → [À rédiger].
    Si c'est le sommaire → inject TOC ou placeholder Word.
    """
    label = PRELIM_LABELS[key]
    _heading(doc, label, level=1)

    if key == 'sommaire':
        if toc['type']:
            _insert_toc(doc, toc)
        else:
            _para(doc,
                "[La table des matières sera générée automatiquement par Word : "
                "Références → Table des matières]",
                size=FONT_BODY, italic=True, align="left")
    elif content_paras:
        _write_section_content(doc, content_paras)
    else:
        _para(doc, "[À rédiger]", size=FONT_BODY, italic=True, align="justify")

    _page_break(doc)


def _write_body_section(doc, paragraphs):
    """Écrit le corps principal du document."""
    for p in paragraphs:
        text = p.text.strip()
        if not text:
            continue
        is_h, level = _is_heading(p)
        if is_h:
            _heading(doc, text, level=level)
        elif re.match(r'^[-•–]\s+', text):
            body_p = doc.add_paragraph(style="List Bullet")
            body_p.paragraph_format.line_spacing = LINE_SPACING
            body_p.paragraph_format.space_after  = Pt(4)
            r = body_p.add_run(re.sub(r'^[-•–]\s+', '', text))
            r.font.name = FONT_NAME
            r.font.size = Pt(FONT_BODY)
        else:
            _para(doc, text, size=FONT_BODY, align="justify",
                  space_before=0, space_after=6)


def _cover_memoire(doc, data):
    color = get_filiere_color(data.get("filiere", "TIC"))
    _para(doc, "UNIVERSITÉ D'ÉBOLOWA", size=13, bold=True, align="center", space_after=2)
    _para(doc, "Faculté des Sciences (FS)", size=12, bold=True, align="center", space_after=2)
    _para(doc, data.get("departement", "Département d'Informatique"), size=11, bold=True, align="center", space_after=2)
    if data.get("laboratoire"):
        _para(doc, data["laboratoire"], size=10, align="center", space_after=2)
    doc.add_paragraph()
    _para(doc, "MÉMOIRE", size=16, bold=True, align="center", color=color, space_after=4)
    _para(doc, "présenté en vue de l'obtention du", size=11, italic=True, align="center", space_after=2)
    d = data.get("diplome", "")
    if data.get("option"):     d += f" – Option : {data['option']}"
    if data.get("specialite"): d += f" – Spécialité : {data['specialite']}"
    _para(doc, d, size=12, bold=True, align="center", space_after=2)
    doc.add_paragraph()
    _para(doc, "Thème :", size=11, bold=True, align="center", space_after=2)
    _para(doc, data.get("theme", ""), size=13, bold=True, align="center", color=color, space_after=12)
    doc.add_paragraph()
    _para(doc, "Présenté par :", size=11, bold=True, align="left", space_after=2)
    _para(doc, data.get("etudiant", ""), size=12, align="left", space_after=2)
    if data.get("matricule"):      _para(doc, f"Matricule : {data['matricule']}", size=11, align="left", space_after=2)
    if data.get("diplome_entree"): _para(doc, f"Diplôme d'entrée : {data['diplome_entree']}", size=11, align="left", space_after=2)
    doc.add_paragraph()
    _para(doc, "Sous la direction de :", size=11, bold=True, align="left", space_after=2)
    enc = data.get("encadreur", "")
    if data.get("grade_encadreur"): enc += f", {data['grade_encadreur']}"
    _para(doc, enc, size=11, align="left", space_after=2)
    if data.get("universite_attache"): _para(doc, data["universite_attache"], size=11, align="left", space_after=2)
    doc.add_paragraph()
    _para(doc, f"Année académique : {data.get('annee_academique', '20XX-20XX')}", size=11, bold=True, align="center")


def _cover_stage(doc, data):
    _para(doc, "UNIVERSITÉ D'ÉBOLOWA", size=13, bold=True, align="center", space_after=2)
    _para(doc, "Faculté des Sciences", size=12, bold=True, align="center", space_after=2)
    if data.get("departement"): _para(doc, data["departement"], size=11, bold=True, align="center", space_after=2)
    doc.add_paragraph()
    _para(doc, "RAPPORT DE STAGE", size=16, bold=True, align="center", space_after=4)
    _para(doc, f"Entreprise : {data.get('entreprise', '')}", size=12, bold=True, align="center", space_after=2)
    _para(doc, f"Période : {data.get('periode', '')}", size=11, align="center", space_after=2)
    doc.add_paragraph()
    _para(doc, "Présenté par :", size=11, bold=True, align="left", space_after=2)
    _para(doc, data.get("etudiant", ""), size=12, align="left", space_after=2)
    if data.get("matricule"): _para(doc, f"Matricule : {data['matricule']}", size=11, align="left", space_after=2)
    doc.add_paragraph()
    if data.get("encadreur_academique"):    _para(doc, f"Encadreur académique : {data['encadreur_academique']}", size=11, align="left", space_after=2)
    if data.get("encadreur_professionnel"): _para(doc, f"Encadreur professionnel : {data['encadreur_professionnel']}", size=11, align="left", space_after=2)
    doc.add_paragraph()
    _para(doc, f"Année académique : {data.get('annee_academique', '')}", size=11, bold=True, align="center")


def _cover_projet(doc, data):
    _para(doc, "UNIVERSITÉ D'ÉBOLOWA", size=13, bold=True, align="center", space_after=2)
    _para(doc, "Faculté des Sciences", size=12, bold=True, align="center", space_after=2)
    if data.get("departement"): _para(doc, data["departement"], size=11, bold=True, align="center", space_after=2)
    doc.add_paragraph()
    _para(doc, "RAPPORT DE PROJET", size=16, bold=True, align="center", space_after=4)
    _para(doc, "Thème :", size=11, bold=True, align="center", space_after=2)
    _para(doc, data.get("theme", ""), size=13, bold=True, align="center", space_after=12)
    doc.add_paragraph()
    _para(doc, "Présenté par :", size=11, bold=True, align="left", space_after=2)
    _para(doc, data.get("etudiant", ""), size=12, align="left", space_after=2)
    if data.get("matricule"): _para(doc, f"Matricule : {data['matricule']}", size=11, align="left", space_after=2)
    doc.add_paragraph()
    if data.get("encadreur_academique"):    _para(doc, f"Encadreur académique : {data['encadreur_academique']}", size=11, align="left", space_after=2)
    if data.get("encadreur_professionnel"): _para(doc, f"Encadreur professionnel : {data['encadreur_professionnel']}", size=11, align="left", space_after=2)
    _para(doc, f"Année académique : {data.get('annee_academique', '')}", size=11, bold=True, align="center")


def _cover_expose(doc, data):
    _para(doc, "UNIVERSITÉ D'ÉBOLOWA", size=13, bold=True, align="center", space_after=2)
    _para(doc, "Faculté des Sciences", size=12, bold=True, align="center", space_after=2)
    if data.get("ue"): _para(doc, f"UE : {data['ue']}", size=11, align="center", space_after=2)
    doc.add_paragraph()
    _para(doc, "EXPOSÉ", size=16, bold=True, align="center", space_after=4)
    _para(doc, "Thème :", size=11, bold=True, align="center", space_after=2)
    _para(doc, data.get("theme", ""), size=13, bold=True, align="center", space_after=12)
    doc.add_paragraph()
    if data.get("membres"):
        _para(doc, f"Groupe {data.get('groupe', '')}", size=11, bold=True, align="left", space_after=2)
        for m in (data["membres"] if isinstance(data["membres"], list) else []):
            if isinstance(m, dict):
                line = m.get("nom", "")
                if m.get("matricule"): line += f"  –  {m['matricule']}"
                _para(doc, line, size=11, align="left", space_after=2)
    doc.add_paragraph()
    if data.get("enseignant"): _para(doc, f"Enseignant : {data['enseignant']}", size=11, align="left", space_after=2)
    if data.get("niveau"):     _para(doc, f"Niveau : {data['niveau']}  |  Filière : {data.get('filiere', '')}", size=11, align="left", space_after=2)
    _para(doc, f"Année académique : {data.get('annee_academique', '')}", size=11, bold=True, align="center")


def _cover_demande(doc, data):
    _para(doc, data.get("etudiant", ""), size=12, bold=True, align="left", space_after=2)
    if data.get("matricule"): _para(doc, f"Matricule : {data['matricule']}", size=11, align="left", space_after=2)
    if data.get("filiere"):   _para(doc, f"Filière : {data['filiere']}  –  Niveau : {data.get('niveau', '')}", size=11, align="left", space_after=2)
    if data.get("email"):     _para(doc, f"Email : {data['email']}", size=11, align="left", space_after=2)
    if data.get("tel"):       _para(doc, f"Tél : {data['tel']}", size=11, align="left", space_after=2)
    doc.add_paragraph()
    _para(doc, data.get("date_lieu", ""), size=11, align="right", space_after=2)
    doc.add_paragraph()
    _para(doc, f"À : {data.get('destinataire', '')}", size=12, bold=True, align="right", space_after=2)
    doc.add_paragraph()
    _para(doc, f"Objet : {data.get('objet', '')}", size=12, bold=True, align="left", space_after=6)
    doc.add_paragraph()
    if data.get("formule_appel"):
        _para(doc, data["formule_appel"], size=12, align="left", space_after=6)
    if data.get("corps"):
        for line in data["corps"].split("\n"):
            if line.strip(): _para(doc, line, size=12, align="justify", space_after=4)
    doc.add_paragraph()
    if data.get("formule_politesse"):
        for line in data["formule_politesse"].split("\n"):
            if line.strip(): _para(doc, line, size=12, align="justify", space_after=4)
    doc.add_paragraph()
    if data.get("pieces_jointes"):
        _para(doc, "Pièces jointes :", size=11, bold=True, align="left", space_after=2)
        for line in data["pieces_jointes"].split("\n"):
            if line.strip(): _para(doc, f"– {line.strip()}", size=11, align="left", space_after=2)
    doc.add_paragraph()
    _para(doc, data.get("signature", data.get("etudiant", "")), size=12, align="right")


COVER_BUILDERS = {
    "memoire":        _cover_memoire,
    "rapport_stage":  _cover_stage,
    "rapport_projet": _cover_projet,
    "expose":         _cover_expose,
    "demande":        _cover_demande,
}


def _build_cover(doc, form_data, doc_type):
    if LOGO_PATH.exists():
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.5))
        except Exception:
            pass
    COVER_BUILDERS.get(doc_type, _cover_expose)(doc, form_data)


def format_document(
    analysis: DocumentAnalysis,
    form_data: dict[str, Any],
    original_text: str,
    output_path: Path,
    source_path: Path | None = None,
) -> Path:
    doc = Document()
    _set_margins(doc.sections[0])
    _apply_base_style(doc)

    doc_type = analysis.document_type

    if doc_type == "demande":
        _cover_demande(doc, form_data)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return output_path

    _build_cover(doc, form_data, doc_type)
    _page_break(doc)

    toc = _extract_toc(source_path) if source_path else {'type': None, 'elements': []}

    if source_path and source_path.exists():
        src_doc   = Document(source_path)
        cover_end = _detect_cover_end(src_doc.paragraphs)
        sections  = _scan_present_sections(src_doc.paragraphs, cover_end)
    else:
        sections = {k: [] for k, _ in PRELIM_SECTIONS}
        sections['body'] = []

    required = REQUIRED_PRELIMS.get(doc_type, [])
    for key in required:
        _write_prelim_section(doc, key, sections.get(key, []), toc, doc_type)

    _write_body_section(doc, sections.get('body', []))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path