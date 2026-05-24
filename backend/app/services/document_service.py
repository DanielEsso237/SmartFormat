import shutil
from pathlib import Path
from fastapi import UploadFile
from app.schemas.document import DocumentUploadResponse
from groq import Groq
import json
from docx import Document
import pdfplumber

class DocumentService:
    def __init__(self):
        self.client = Groq()
        self.charte_dir = Path("app/charte")
        
        with open(self.charte_dir / "formatting.json", encoding="utf-8") as f:
            self.formatting_rules = json.load(f)
        with open(self.charte_dir / "academic_rules.json", encoding="utf-8") as f:
            self.academic_rules = json.load(f)
        with open(self.charte_dir / "advanced_formatting.json", encoding="utf-8") as f:
            self.advanced_formatting = json.load(f)

    async def process_document(self, file: UploadFile) -> DocumentUploadResponse:
        temp_path = Path(f"temp_{file.filename}")
        with open(temp_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        try:
   
            if file.filename.endswith(".pdf"):
                text = self._extract_text_from_pdf(temp_path)
            else:
                text = self._extract_text_from_docx(temp_path)
                
            analysis = await self._analyze_with_ai(text, file.filename)

            doc_type = analysis.get("document_type", "memoire")

            output_filename = f"formatted_{file.filename.replace('.pdf', '.docx')}"
            output_path = Path("uploads") / output_filename

 
            Path("uploads").mkdir(exist_ok=True)

     
            self._create_basic_formatted_doc(text, output_path, doc_type)

            return DocumentUploadResponse(
                filename=output_filename,
                original_name=file.filename,
                file_type="docx",
                status="success",
                message="Document analysé et formaté avec succès",
                download_url=f"/api/download/{output_filename}",
                document_type=doc_type
            )

        finally:
            if temp_path.exists():
                temp_path.unlink()

    def _extract_text_from_pdf(self, path: Path) -> str:
        text = ""
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text

    def _extract_text_from_docx(self, path: Path) -> str:
        doc = Document(path)
        return "\n".join([para.text for para in doc.paragraphs])

    async def _analyze_with_ai(self, text: str, filename: str):
        prompt = f"""Tu es un expert en mise en forme académique selon la charte de la Faculté des Sciences.
Analyse ce document et retourne uniquement un JSON.

Document: {text[:8000]}...

Retourne ce JSON :
{{
  "document_type": "memoire | rapport_stage | rapport_projet | expose | demande",
  "title": "...",
  "has_cover": true/false,
  "needs_form": true/false
}}
"""

        response = self.client.chat.completions.create(
            model="llama-3.1-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=500
        )

        try:
            return json.loads(response.choices[0].message.content)
        except:
            return {"document_type": "memoire", "title": "Document Académique", "has_cover": False, "needs_form": True}

    def _create_basic_formatted_doc(self, text: str, output_path: Path, doc_type: str):
        doc = Document()
        
        style = doc.styles['Normal']
        font = style.font
        font.name = self.formatting_rules["font"]["name"]
        font.size = self.formatting_rules["font"]["size"] * 0.5  

        doc.add_paragraph(text[:2000]) 
        
        doc.save(output_path)